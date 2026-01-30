"""Generate DOCX risk assessment documents."""

from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import RiskAssessment


def get_risk_color(risk_level: str) -> RGBColor:
    """Get color for risk level."""
    colors = {
        "Low": RGBColor(40, 167, 69),      # Green
        "Medium": RGBColor(255, 193, 7),   # Amber
        "High": RGBColor(220, 53, 69),     # Red
    }
    return colors.get(risk_level, RGBColor(0, 0, 0))


def generate_docx(assessment: RiskAssessment) -> BytesIO:
    """Generate a DOCX document from a risk assessment.

    Args:
        assessment: The completed risk assessment

    Returns:
        BytesIO buffer containing the DOCX file
    """
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Risk Assessment', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle with activity name
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(assessment.activity_name)
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(16)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Activity Details Table
    doc.add_heading('Activity Details', level=1)

    details_table = doc.add_table(rows=6, cols=2)
    details_table.style = 'Table Grid'

    details = [
        ('Activity Name', assessment.activity_name),
        ('Description', assessment.activity_description),
        ('Location', assessment.location),
        ('Age Groups', ', '.join(ag.value for ag in assessment.age_groups)),
        ('Assessment Date', assessment.assessment_date.strftime('%d %B %Y')),
        ('Assessed By', assessment.assessor_name or 'Not specified'),
    ]

    for i, (label, value) in enumerate(details):
        row = details_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value

    # Set column widths
    for row in details_table.rows:
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(4.5)

    doc.add_paragraph()

    # Overall Risk Level
    risk_para = doc.add_paragraph()
    risk_para.add_run('Overall Risk Level: ').bold = True
    risk_run = risk_para.add_run(assessment.overall_risk_level)
    risk_run.bold = True
    risk_run.font.color.rgb = get_risk_color(assessment.overall_risk_level)

    doc.add_paragraph()

    # Hazards Table
    doc.add_heading('Identified Hazards & Control Measures', level=1)

    for i, hwm in enumerate(assessment.hazards, 1):
        h = hwm.hazard

        # Hazard heading
        hazard_heading = doc.add_paragraph()
        hazard_heading.add_run(f'{i}. {h.description}').bold = True

        # Risk info
        risk_info = doc.add_paragraph()
        risk_info.add_run('Severity: ').bold = True
        risk_info.add_run(f'{h.severity.value}  |  ')
        risk_info.add_run('Likelihood: ').bold = True
        risk_info.add_run(f'{h.likelihood.value}  |  ')
        risk_info.add_run('Risk Level: ').bold = True
        level_run = risk_info.add_run(h.risk_level)
        level_run.font.color.rgb = get_risk_color(h.risk_level)

        # Who at risk
        who_para = doc.add_paragraph()
        who_para.add_run('Who is at risk: ').bold = True
        who_para.add_run(h.who_at_risk)

        # Existing controls
        if hwm.existing_controls:
            controls_para = doc.add_paragraph()
            controls_para.add_run('Existing Controls:').bold = True
            for control in hwm.existing_controls:
                doc.add_paragraph(control, style='List Bullet')

        # Additional controls
        if hwm.additional_controls:
            additional_para = doc.add_paragraph()
            additional_para.add_run('Additional Controls Required:').bold = True
            for control in hwm.additional_controls:
                doc.add_paragraph(
                    f'{control.action} ({control.responsible_person})',
                    style='List Bullet'
                )

        # Residual risk
        residual_para = doc.add_paragraph()
        residual_para.add_run('Residual Risk: ').bold = True
        residual_run = residual_para.add_run(hwm.residual_risk)
        residual_run.font.color.rgb = get_risk_color(hwm.residual_risk)

        doc.add_paragraph()  # Spacing between hazards

    # Additional Notes
    if assessment.additional_notes:
        doc.add_heading('Additional Notes', level=1)
        doc.add_paragraph(assessment.additional_notes)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('This risk assessment was generated in accordance with the ')
    footer.add_run('Statutory Framework for the Early Years Foundation Stage (EYFS) 2024').italic = True
    footer.add_run('. Review date: ')
    if assessment.review_date:
        footer.add_run(assessment.review_date.strftime('%d %B %Y'))
    else:
        footer.add_run('To be determined')

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
